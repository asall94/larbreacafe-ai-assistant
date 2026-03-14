#!/usr/bin/env python3
"""Complete simulation of the L'Arbre à Café chatbot."""

import requests
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    # Create a hyperlink relationship
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the link text
    new_run = OxmlElement('w:r')

    # Set run properties (link style)
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font Calibri 11
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt = 21 half-points
    rPr.append(sz)

    new_run.append(rPr)

    # Add the text
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def run_simulation():
    questions = [
        'Proposez-vous des cafés de la ferme Mariposa ?',
        'Quel est le prix du café Source en 250g ?',
        'Avez-vous une boutique dans le 9ème arrondissement ?',
        'Combien de boutiques avez-vous à Paris ?',
        'Proposez-vous des formations pour devenir barista ?',
        'Vendez-vous des machines Sage The Barista Pro ?',
        'Proposez-vous des cafés biodynamiques ?',
        'Avez-vous des cafés d\'Éthiopie ?',
        'Quels sont vos cafés d\'exception ?',
        'Comment contacter la boutique Paris 09 ?'
    ]
    
    results = []
    
    print('\n' + '='*70)
    print('   SIMULATION CHATBOT L\'ARBRE À CAFÉ - AGENTIC RAG')
    print(f'   Date: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    print('='*70)
    
    for i, q in enumerate(questions, 1):
        print(f'\n[Q{i}] {q}')
        print('-'*70)
        try:
            r = requests.post('http://localhost:8002/chat', 
                            json={'message': q}, 
                            timeout=30)
            response = r.json().get('response', 'Erreur')
            print(response)
            results.append({
                'numero': i,
                'question': q, 
                'response': response,
                'status': 'ok'
            })
        except Exception as e:
            error_msg = f'ERREUR: {str(e)}'
            print(error_msg)
            results.append({
                'numero': i,
                'question': q, 
                'response': error_msg,
                'status': 'error'
            })
    
    # Save results to JSON with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    json_filename = f'simulation_results_{timestamp}.json'
    output_data = {
        'date': datetime.now().isoformat(),
        'company': 'L\'Arbre à Café',
        'total_questions': len(questions),
        'questions': results
    }
    
    with open(json_filename, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    # Export results to DOCX
    docx_filename = f'simulation_results_{timestamp}.docx'
    doc = Document()
    
    # Main title - centered and underlined
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Agent intelligent pour L'Arbre à Café - Simulation du {datetime.now().strftime('%d/%m/%Y')}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.underline = True
    
    doc.add_paragraph()  # Empty line
    
    # Add each question/response pair
    for r in results:
        # Question text
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f"Q{r['numero']} : {r['question']}")
        question_run.font.name = 'Calibri'
        question_run.font.size = Pt(10.5)
        
        doc.add_paragraph()  # Empty line
        
        # Response - process HTML links
        response_text = r['response']
        response_para = doc.add_paragraph()
        
        # Parse and add text with hyperlinks
        # Pattern to extract links: <a href="URL" target="_blank">Text</a>
        link_pattern = r'<a href="([^"]+)"[^>]*>([^<]+)</a>'
        last_pos = 0
        
        for match in re.finditer(link_pattern, response_text):
            # Text before the link
            before_text = response_text[last_pos:match.start()]
            if before_text:
                run = response_para.add_run(before_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
            
            # Add the clickable link using add_hyperlink
            url = match.group(1)
            link_text = match.group(2)
            add_hyperlink(response_para, url, link_text)
            
            last_pos = match.end()
        
        # Text after the last link
        remaining_text = response_text[last_pos:]
        if remaining_text:
            # Strip any remaining HTML tags
            remaining_text = re.sub(r'<[^>]+>', '', remaining_text)
            run = response_para.add_run(remaining_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
                
        # Separator
        separator = doc.add_paragraph('---')
        sep_run = separator.runs[0]
        sep_run.font.name = 'Calibri'
        sep_run.font.size = Pt(10.5)
        
        doc.add_paragraph()  # Empty line
    
    doc.save(docx_filename)
    
    print('\n' + '='*70)
    print(f'[OK] JSON: {json_filename}')
    print(f'[OK] DOCX: {docx_filename}')
    print('='*70)
    
    return results

if __name__ == "__main__":
    run_simulation()
